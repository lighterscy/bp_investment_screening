"""Formal Chinese-style Word report rendering."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from bp_investment_screening.llm import LLMClient
from bp_investment_screening.report_composer import ReportSectionComposer
from bp_investment_screening.schemas import EvidenceItem, InvestmentMemo, Layer1ResearchItem
from bp_investment_screening.technical_background import generate_technical_background_sections


TITLE_FONT = "方正小标宋简体"
BODY_FONT = "仿宋_GB2312"
HEADING_FONT = "黑体"
SECOND_HEADING_FONT = "楷体_GB2312"
FALLBACK_EAST_ASIA_FONT = "宋体"


class FormalDocxWriter:
    """Render the final project screening memo as a formal Word document."""

    def __init__(self, llm_client: LLMClient | None = None) -> None:
        self.llm_client = llm_client

    def write(self, memo: InvestmentMemo, output_dir: str | Path) -> Path:
        root = Path(output_dir)
        root.mkdir(parents=True, exist_ok=True)
        doc = Document()
        _setup_document(doc)
        _render_document(doc, memo, self.llm_client)
        path = root / "project_research_memo.docx"
        doc.save(path)
        return path


def _setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(16)
    _set_east_asia_font(normal, BODY_FONT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(28)
    normal.paragraph_format.first_line_indent = Pt(32)


def _render_document(
    doc: Document,
    memo: InvestmentMemo,
    llm_client: LLMClient | None,
) -> None:
    _add_title(doc, f"{memo.project_name}初步研判报告")
    composer = ReportSectionComposer(llm_client=llm_client)
    _add_first_heading(doc, "一、项目基本情况梳理")
    _add_second_heading(doc, "（一）公司基础概况")
    _add_body(doc, composer.compose(memo, "公司基础概况", _company_overview(memo), use_llm=False))
    _add_second_heading(doc, "（二）股权结构与融资计划")
    _add_body(doc, composer.compose(memo, "股权结构与融资计划", "BP 未稳定披露股权结构与融资计划，需在下一步 DD 中补充确认。", use_llm=False))
    _add_second_heading(doc, "（三）核心团队情况")
    _add_body(doc, composer.compose(memo, "核心团队情况", "BP 未稳定披露核心团队情况，需补充创始人履历、组织分工与关键技术负责人背景。", use_llm=False))
    _add_second_heading(doc, "（四）技术背景")
    _add_body(doc, "本部分用于说明项目涉及的技术门类本身，帮助读者先理解基础概念、上位体系、关键原理和技术演进逻辑；具体公司技术进展在后文“产品与技术情况”中展开。")
    for section in generate_technical_background_sections(memo, llm_client=llm_client):
        _add_third_heading(doc, section.heading)
        _add_body(doc, section.body)

    _add_second_heading(doc, "（五）产品与技术情况")
    _add_third_heading(doc, "1、技术情况")
    _add_body(doc, composer.compose(memo, "技术情况", "项目技术情况待补充。", use_llm=False))
    _add_third_heading(doc, "2、产品情况")
    _add_body(doc, composer.compose(memo, "产品情况", _product_text(memo), use_llm=False))
    _add_third_heading(doc, "3、知产情况")
    _add_body(doc, composer.compose(memo, "知产情况", "BP 中如披露专利、软件著作权、核心工艺或技术秘密，应在此处列示；当前需进一步核验知识产权权属、有效性与保护边界。", use_llm=False))

    _add_second_heading(doc, "（六）业务与市场布局")
    _add_body(doc, composer.compose(memo, "业务与市场布局", "业务、客户、订单及市场布局信息待补充核验。", use_llm=False))
    _add_second_heading(doc, "（七）业绩及发展规划")
    _add_body(doc, composer.compose(memo, "业绩及发展规划", "业绩数据、收入构成、毛利水平、订单可实现性及未来规划均应作为待验证事项，后续需调取财务报表、合同/订单与客户访谈材料。", use_llm=False))

    _add_first_heading(doc, "二、项目所处行业基本情况分析")
    _add_second_heading(doc, "（一）行业核心定义")
    _add_body(doc, composer.compose(memo, "行业核心定义", _industry_definition(memo), use_llm=False))
    _add_second_heading(doc, "（二）行业发展现状")
    _add_third_heading(doc, "1、海外发展现状")
    _add_body(doc, composer.compose(memo, "海外发展现状", "待通过外部资料补充海外行业阶段、主要企业、技术路线及产业化节奏。", use_llm=False))
    _add_third_heading(doc, "2、国内发展现状")
    _add_body(doc, composer.compose(memo, "国内发展现状", "待通过外部资料补充国内政策环境、产业链成熟度、关键供给缺口及国产替代进展。", use_llm=False))
    _add_third_heading(doc, "3、技术突破与核心痛点")
    _add_body(doc, composer.compose(memo, "技术突破与核心痛点", "行业技术突破与核心痛点待外部证据补充。", use_llm=False))
    _add_second_heading(doc, "（三）行业市场规模")
    _add_third_heading(doc, "1、国内市场规模")
    _add_body(doc, composer.compose(memo, "国内市场规模", "待通过可信第三方资料或产业链访谈补充国内市场规模、增速及可服务市场范围。", use_llm=False))
    _add_third_heading(doc, "2、全球市场规模")
    _add_body(doc, composer.compose(memo, "全球市场规模", "待通过可信第三方资料补充全球市场规模、主要区域分布及国际竞争格局。", use_llm=False))
    _add_second_heading(doc, "（四）行业主要竞争对手")
    _add_body(doc, composer.compose(memo, "行业主要竞争对手", "主要竞争对手、替代方案及项目差异化待外部证据补充。", use_llm=False))

    _add_first_heading(doc, "三、初步建议")
    _add_second_heading(doc, "（一）项目优势")
    _add_body(doc, _list_to_sentence(memo.recommendation.strengths))
    _add_second_heading(doc, "（二）项目疑问点")
    doubts = memo.recommendation.weaknesses + memo.recommendation.key_risks
    _add_body(doc, _list_to_sentence(doubts))
    _add_second_heading(doc, "（三）初步判断")
    _add_body(
        doc,
        f"基于当前 BP 信息和已完成的初筛归集，项目初筛建议为“{memo.recommendation.recommendation}”，"
        f"判断置信度为“{_confidence_label(memo.recommendation.confidence)}”。{memo.recommendation.one_sentence_view}",
    )


def _add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(text)
    _set_run_font(run, TITLE_FONT, Pt(22), bold=False)


def _add_first_heading(doc: Document, text: str) -> None:
    _add_heading(doc, text, HEADING_FONT, Pt(16), bold=True)


def _add_second_heading(doc: Document, text: str) -> None:
    _add_heading(doc, text, SECOND_HEADING_FONT, Pt(16), bold=True)


def _add_third_heading(doc: Document, text: str) -> None:
    _add_heading(doc, text, BODY_FONT, Pt(16), bold=True)


def _add_heading(doc: Document, text: str, font: str, size: Pt, bold: bool) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(28)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    _set_run_font(run, font, size, bold=bold)


def _add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(32)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(28)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    _set_run_font(run, BODY_FONT, Pt(16), bold=False)


def _set_run_font(run, font_name: str, size: Pt, bold: bool) -> None:
    run.bold = bold
    run.font.name = font_name
    run.font.size = size
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def _set_east_asia_font(style, font_name: str) -> None:
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def _company_overview(memo: InvestmentMemo) -> str:
    return (
        f"{memo.project_name}所属行业为{memo.industry or '待确认'}。"
        "公司基础信息主要来源于 BP，当前阶段仅作为项目方披露信息归集，后续需通过工商信息、访谈及底层材料进一步核验。"
    )


def _product_text(memo: InvestmentMemo) -> str:
    product_item = _find_item(memo, "产品与服务")
    if product_item and product_item.bp_claims:
        return _evidence_sentence(product_item.bp_claims)
    return "BP 未稳定披露产品形态与服务边界，需进一步明确核心产品、应用场景、交付方式与客户使用效果。"


def _industry_definition(memo: InvestmentMemo) -> str:
    return (
        f"本项目暂按{memo.industry or '待确认行业'}进行初步归类。"
        "行业边界、上游供给、下游应用及项目所处产业链环节需结合外部研究进一步界定。"
    )


def _claims_text(memo: InvestmentMemo, topics: tuple[str, ...], fallback: str) -> str:
    claims: list[EvidenceItem] = []
    for topic in topics:
        item = _find_item(memo, topic)
        if item:
            claims.extend(item.bp_claims[:4])
    if claims:
        return _evidence_sentence(claims)
    return fallback


def _find_item(memo: InvestmentMemo, topic: str) -> Layer1ResearchItem | None:
    for item in memo.layer1_items:
        if item.topic == topic:
            return item
    return None


def _evidence_sentence(evidence: list[EvidenceItem]) -> str:
    parts = []
    for item in evidence[:4]:
        page = f"（BP第{item.source_page}页）" if item.source_page else ""
        parts.append(f"{item.content}{page}")
    return "；".join(parts) + "。上述信息均需在后续 DD 中进一步核验。"


def _list_to_sentence(items: list[str]) -> str:
    if not items:
        return "当前证据不足，需在下一步 DD 中补充核验。"
    return "；".join(items) + "。"


def _confidence_label(confidence: str) -> str:
    return {
        "low": "低",
        "medium": "中",
        "high": "高",
    }.get(confidence, confidence)
